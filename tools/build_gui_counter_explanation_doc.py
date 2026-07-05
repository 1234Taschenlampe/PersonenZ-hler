from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOT = Path(r"C:\Users\Felix\AppData\Local\Temp\codex-clipboard-ca27988d-069f-4bb4-bc20-f408e1d457ca.png")
OUT_DIR = ROOT / "artifacts" / "gui_counter_explanation"
ANNOTATED = OUT_DIR / "gui_counter_markiert.png"
DOCX = OUT_DIR / "GUI_Counter_Erklaerung.docx"


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def annotate_screenshot() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.open(SCREENSHOT).convert("RGB")
    draw = ImageDraw.Draw(image)

    try:
        font = ImageFont.truetype("arial.ttf", 24)
        small = ImageFont.truetype("arial.ttf", 18)
    except OSError:
        font = ImageFont.load_default()
        small = ImageFont.load_default()

    def label_box(xy, label: str, color: tuple[int, int, int], fill: tuple[int, int, int]) -> None:
        x1, y1, x2, y2 = xy
        draw.rectangle(xy, outline=color, width=5)
        text_box = draw.textbbox((0, 0), label, font=font)
        tw = text_box[2] - text_box[0]
        th = text_box[3] - text_box[1]
        draw.rectangle((x1, max(0, y1 - th - 14), x1 + tw + 18, y1), fill=fill, outline=color, width=2)
        draw.text((x1 + 9, max(2, y1 - th - 10)), label, fill=(255, 255, 255), font=font)

    label_box((125, 650, 380, 758), "ROT: Global Counter pruefen", (210, 0, 0), (210, 0, 0))
    label_box((125, 758, 380, 930), "ORANGE: Kamera-Zaehler Vergleich", (230, 120, 0), (230, 120, 0))
    label_box((783, 648, 1620, 1018), "BLAU: Status / Hailo / Latenz", (0, 86, 170), (0, 86, 170))
    label_box((1625, 650, 1792, 1020), "GRUEN: Steuerung", (0, 130, 70), (0, 130, 70))
    label_box((386, 650, 780, 1018), "VIOLETT: Kamera-Auswahl", (120, 60, 170), (120, 60, 170))
    label_box((120, 155, 1796, 633), "Kamerabilder + Zaehllinien", (0, 120, 120), (0, 120, 120))

    image.save(ANNOTATED)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "YOLO26m Besucherzaehler - GUI und Counter-Erklaerung"
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(90, 90, 90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Hinweis: Rot markierte Punkte sind die Stellen, an denen der globale Counter geprueft werden muss."
    fp.runs[0].font.name = "Arial"
    fp.runs[0].font.size = Pt(8)
    fp.runs[0].font.color.rgb = RGBColor(120, 120, 120)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_status_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Bereich", "Was bedeutet es?", "Aktueller Eindruck", "Markierung"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_fill(table.rows[0].cells[i], "EDEDED")

    rows = [
        (
            "Global inside / global in / global out",
            "Gesamtzaehler fuer alle Kameras. Diese Werte sollen echte Ein- und Ausgaenge zusammenfassen.",
            "Auf dem Screenshot bleiben alle drei Werte 0. Das ist die wichtigste Problemstelle.",
            "ROT",
        ),
        (
            "camera 1 visible / camera 2 visible",
            "Zeigt, ob aktuell Personen/Tracks im Bild sichtbar sind. Das ist nicht automatisch ein Eintritt.",
            "camera 1 visible zeigt 1, globale Werte bleiben aber 0.",
            "ORANGE",
        ),
        (
            "camera 1 in/out und camera 2 in/out",
            "Lokale Linienueberquerungen pro Kamera.",
            "Wenn diese Werte steigen, globale Werte aber nicht, liegt der Fehler nach der lokalen Erkennung.",
            "ORANGE",
        ),
        (
            "Status: backend, hailo, hef, osnet",
            "Zeigt, ob YOLO26m/Hailo und ReID aktiv sind.",
            "YOLO26m und Hailo sind aktiv. Das Modell ist nicht der offensichtliche Fehler.",
            "BLAU",
        ),
        (
            "Zaehllinie im Bild",
            "Die gelbe Linie ist die Grenze. Gezaehlt wird erst bei einer stabil erkannten Ueberquerung.",
            "Wenn die Linie falsch liegt oder Richtung falsch ist, zaehlt global nichts.",
            "TUERKIS",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
        if row[3] == "ROT":
            for cell in cells:
                set_cell_fill(cell, "F4CCCC")
        elif row[3] == "ORANGE":
            for cell in cells:
                set_cell_fill(cell, "FCE5CD")
        elif row[3] == "BLAU":
            for cell in cells:
                set_cell_fill(cell, "D9EAF7")
        else:
            for cell in cells:
                set_cell_fill(cell, "D9EAD3")


def build_doc() -> None:
    annotate_screenshot()

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Arial"

    add_header_footer(doc)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("GUI-Erklaerung und Counter-Pruefpunkte")
    r.font.color.rgb = RGBColor(45, 65, 90)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Stand: {datetime.now().strftime('%d.%m.%Y %H:%M')} | System: YOLO26m Dual-Kamera Besucherzaehler").italic = True

    p = doc.add_paragraph()
    p.add_run("Zweck: ").bold = True
    p.add_run(
        "Dieses Dokument erklaert die sichtbaren GUI-Bereiche und markiert die Stellen, "
        "an denen der globale Counter geprueft werden muss."
    )

    doc.add_heading("1. Markierter Screenshot", level=1)
    doc.add_picture(str(ANNOTATED), width=Inches(8.8))
    cap = doc.add_paragraph("Rot = globale Counter-Problemstelle. Orange = Kamera-Werte zum Vergleich. Blau/Gruen/Violett = Status, Steuerung und Kamera-Auswahl.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("2. Was ist was?", level=1)
    add_status_table(doc)

    doc.add_heading("3. Globaler Counter: was aktuell nicht passt", level=1)
    issue_table = doc.add_table(rows=1, cols=3)
    issue_table.style = "Table Grid"
    for i, header in enumerate(["Markierung", "Beobachtung", "Was pruefen?"]):
        set_cell_text(issue_table.rows[0].cells[i], header, bold=True)
        set_cell_fill(issue_table.rows[0].cells[i], "EDEDED")
    issue_rows = [
        (
            "ROT",
            "global inside, global in und global out bleiben 0.",
            "Pruefen, ob CrossingEvent entsteht und ob EventDatabase.record_decision den globalen Zaehler aktualisiert.",
        ),
        (
            "ROT",
            "Person ist sichtbar, aber global bleibt 0.",
            "Das ist nur dann ein Fehler, wenn die Person die gelbe Linie wirklich in der richtigen Richtung gekreuzt hat.",
        ),
        (
            "ROT",
            "camera visible ist kein global inside.",
            "Falls gewuenscht, muss getrennt entschieden werden: Live-Anwesenheit anzeigen oder echte Ein-/Ausgaenge zaehlen.",
        ),
        (
            "ROT",
            "Wenn camera 1 in/out oder camera 2 in/out steigen, global aber nicht.",
            "Dann liegt der Fehler vermutlich zwischen lokalem Kamera-Counter, Konsenslogik und globaler Datenbank-Aktualisierung.",
        ),
    ]
    for row in issue_rows:
        cells = issue_table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            set_cell_fill(cells[i], "F4CCCC")

    doc.add_page_break()
    doc.add_heading("4. Kurzbedienung", level=1)
    controls = [
        ("Start", "Startet Kameras und Verarbeitung."),
        ("Stopp", "Stoppt Verarbeitung."),
        ("Neustart", "Startet Capture und Pipeline neu."),
        ("Zaehler zuruecksetzen", "Setzt sichtbare Zaehler zurueck."),
        ("Kameras neu erkennen", "Sucht USB-Kameras neu und aktualisiert die Auswahl."),
        ("Auswahl uebernehmen / speichern", "Uebernimmt die Kamera-Zuordnung und schreibt sie in config/config.yaml."),
        ("Diagnosebericht", "Schreibt einen Diagnosebericht in logs/."),
    ]
    ctl_table = doc.add_table(rows=1, cols=2)
    ctl_table.style = "Table Grid"
    set_cell_text(ctl_table.rows[0].cells[0], "Button", bold=True)
    set_cell_text(ctl_table.rows[0].cells[1], "Funktion", bold=True)
    set_cell_fill(ctl_table.rows[0].cells[0], "EDEDED")
    set_cell_fill(ctl_table.rows[0].cells[1], "EDEDED")
    for button, meaning in controls:
        row = ctl_table.add_row().cells
        set_cell_text(row[0], button, bold=True)
        set_cell_text(row[1], meaning)

    note = doc.add_paragraph()
    note.add_run("Wichtig: ").bold = True
    note.add_run(
        "Der globale Counter zaehlt keine reine Sichtbarkeit. Er muss erst steigen, wenn eine Person stabil erkannt "
        "die gelbe Linie in der konfigurierten Richtung ueberquert. Wenn das passiert und global bleibt 0, ist das der echte Fehlerpfad."
    )

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build_doc()
